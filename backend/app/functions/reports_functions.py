from io import BytesIO
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


def create_employee_report_func(filename, full_name, age, contact_info, role, period, working_days, days_off):
    doc = Document()

    title = doc.add_heading('Отчёт о сотруднике', level=1)
    title.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    title.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    title.paragraph_format.space_after = Pt(12)
    title.alignment = 1

    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 0, 0)

    title1 = doc.add_heading('Основная информация', level=2)

    for run in title1.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("ФИО: ").bold = True
    p.add_run(full_name).font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Должность: ").bold = True
    p.add_run(str(role)).font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Возраст: ").bold = True
    p.add_run(str(age)).font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Контактные данные: ").bold = True
    p.add_run(contact_info).font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет

    title2 = doc.add_heading('Период отчёта', level=2)

    for run in title2.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Период: ").bold = True
    p.add_run(period).font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Количество рабочих дней: ").bold = True
    p.add_run(str(working_days)).font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Количество выходных дней: ").bold = True
    p.add_run(str(days_off)).font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run(
        f"За указанный период сотрудник работал {round(working_days / (working_days + days_off) * 100, 1)}% времени"
    ).bold = True

    fig, ax = plt.subplots()
    categories = ['Рабочие дни', 'Выходные дни']
    values = [working_days, days_off]
    ax.bar(categories, values, color=['darkblue', 'darkblue'])
    ax.set_title('Рабочие и выходные дни')
    ax.set_ylabel('Количество дней')

    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)

    paragraph = doc.add_paragraph()  # Создаем абзац для изображения
    run = paragraph.add_run()  # Добавляем run в абзац
    run.add_picture(buf, width=Inches(4))  # Вставляем изображение
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    plt.close(fig)  # Закрываем фигуру, чтобы освободить память

    doc.save(filename)


def create_point_report_func(filename, point, people_who_wish, people_who_work, period, working_days, days_off):
    doc = Document()

    title = doc.add_heading('Отчёт о точке', level=1)
    title.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    title.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    title.paragraph_format.space_after = Pt(12)
    title.alignment = 1

    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 0, 0)

    title1 = doc.add_heading('Основная информация', level=2)

    for run in title1.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Адрес точки: ").bold = True
    p.add_run(point).font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Сотрудники, у которых данная точка отмечена желаемой: ").bold = True
    p.add_run(', '.join(people_who_wish)).font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет

    title2 = doc.add_heading('Период отчёта', level=2)

    for run in title2.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Период: ").bold = True
    p.add_run(period).font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Сотрудники, которые работали на данной точке в указанный период: ").bold = True
    p.add_run(', '.join(people_who_work)).font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Количество отработанных дней: ").bold = True
    p.add_run(str(working_days)).font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run("Количество пропущенных дней: ").bold = True
    p.add_run(str(days_off)).font.color.rgb = RGBColor(0, 0, 0)  # Чёрный цвет

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5  # Полуторный интервал
    p.paragraph_format.space_before = Pt(12)  # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(12)
    p.add_run(
        f"За указанный период точка отработала {round(working_days / (working_days + days_off) * 100, 1)}% времени"
    ).bold = True

    fig, ax = plt.subplots()
    categories = ['Отработанные дни', 'Пропущенные дни']
    values = [working_days, days_off]
    ax.bar(categories, values, color=['darkblue', 'darkblue'])
    ax.set_title('Отработанные и Пропущенные дни')
    ax.set_ylabel('Количество дней')

    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)

    paragraph = doc.add_paragraph()  # Создаем абзац для изображения
    run = paragraph.add_run()  # Добавляем run в абзац
    run.add_picture(buf, width=Inches(4))  # Вставляем изображение
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    plt.close(fig)  # Закрываем фигуру, чтобы освободить память

    doc.save(filename)
